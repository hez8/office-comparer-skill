import streamlit as st
import pandas as pd
from docx import Document
from PIL import Image
import numpy as np
import difflib
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io
import cv2
import time

# --- 悬浮办公风格设计 ---
STYLE = """
<style>
    .stApp { background-color: #f3f2f1; font-family: 'Segoe UI', sans-serif; }
    div[data-testid="stAppViewBlockContainer"] { max-width: 98% !important; padding-top: 5rem !important; }
    .app-header {
        position: fixed; top: 0; left: 15%; right: 15%; height: 50px;
        background-color: #ffffff; display: flex; align-items: center; justify-content: center;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1); z-index: 999;
        border-radius: 0 0 12px 12px; border: 1px solid #0078d4; border-top: none;
    }
    .app-title { color: #0078d4; font-size: 1.3rem; font-weight: 700; display: flex; align-items: center; gap: 10px; }
    .line-num { color: #888; font-size: 0.8rem; text-align: right; padding-right: 8px; border-right: 1px solid #ddd; min-width: 40px; }
    .doc-content { padding: 6px 12px; font-size: 0.95rem; background-color: white; border-bottom: 1px solid #f3f2f1; line-height: 1.6; }
    .diff-removed { background-color: #fff4ce !important; border-left: 5px solid #ffb900; }
    .diff-added { background-color: #dff6dd !important; border-left: 5px solid #107c10; }
    .char-del { text-decoration: line-through; color: #d13438; background-color: #fde7e9; padding: 0 2px; }
    .char-add { color: #107c10; background-color: #dff6dd; font-weight: bold; padding: 0 2px; }
    .stButton>button { background-color: #0078d4; color: white; border-radius: 4px; border: none; width: 100%; }
</style>
"""

HEADER_HTML = """
<div class="app-header">
    <div class="app-title">
        <svg width="24" height="24" viewBox="0 0 32 32" fill="none"><path d="M18 2H8C6.9 2 6.01 2.9 6.01 4L6 28C6 29.1 6.89 30 7.99 30H24C25.1 30 26 29.1 26 28V10L18 2Z" fill="#0078D4"/><path d="M18 2V10H26L18 2Z" fill="#C7E0F4"/><path d="M12 16H20V18H12V16ZM12 20H20V22H12V20ZM12 12H16V14H12V12Z" fill="white"/></svg>
        Office-Comparer Pro
    </div>
</div>
"""

# --- Logic: Document Processing ---

@st.cache_data
def load_docx_lines(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        return [p.text for p in doc.paragraphs]
    except Exception as e:
        st.error(f"无法读取 Word 文件: {e}")
        return []

def save_docx_from_lines(lines):
    doc = Document()
    for line in lines: doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def highlight_text_diff(text_a, text_b):
    """实现字符级差异高亮"""
    result_a, result_b = [], []
    s = difflib.SequenceMatcher(None, text_a, text_b)
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        if tag == 'equal':
            result_a.append(text_a[i1:i2]); result_b.append(text_b[j1:j2])
        elif tag == 'delete':
            result_a.append(f'<span class="char-del">{text_a[i1:i2]}</span>')
        elif tag == 'insert':
            result_b.append(f'<span class="char-add">{text_b[j1:j2]}</span>')
        elif tag == 'replace':
            result_a.append(f'<span class="char-del">{text_a[i1:i2]}</span>')
            result_b.append(f'<span class="char-add">{text_b[j1:j2]}</span>')
    return "".join(result_a), "".join(result_b)

# --- Logic: Excel Processing ---

@st.cache_data
def load_excel_df(file_bytes):
    try:
        return pd.read_excel(io.BytesIO(file_bytes))
    except Exception as e:
        st.error(f"无法读取 Excel 文件: {e}")
        return None

def compare_dataframes(df_a, df_b):
    """对比两个 DataFrame 并生成高亮样式"""
    # 统一尺寸
    rows = max(len(df_a), len(df_b))
    cols = max(len(df_a.columns), len(df_b.columns))
    
    # 填充缺失
    df_a_ext = df_a.reindex(index=range(rows), columns=df_a.columns.tolist() + [f"empty_{i}" for i in range(cols-len(df_a.columns))]).fillna("")
    df_b_ext = df_b.reindex(index=range(rows), columns=df_b.columns.tolist() + [f"empty_{i}" for i in range(cols-len(df_b.columns))]).fillna("")
    
    def highlight_cells(x):
        df_mask = pd.DataFrame('', index=x.index, columns=x.columns)
        for r in range(rows):
            for c in range(len(x.columns)):
                val_a = str(df_a_ext.iloc[r, c]) if r < len(df_a_ext) else None
                val_b = str(df_b_ext.iloc[r, c]) if r < len(df_b_ext) else None
                if val_a != val_b:
                    df_mask.iloc[r, c] = 'background-color: #fff4ce; border: 1px solid #ffb900'
        return df_mask

    return df_a_ext.style.apply(highlight_cells, axis=None), df_b_ext.style.apply(highlight_cells, axis=None)

# --- Logic: Image Processing ---

@st.cache_data
def align_images_cv2(img_a_bytes, img_b_bytes):
    """使用 ORB 算法进行图像自动对齐"""
    nparr_a = np.frombuffer(img_a_bytes, np.uint8)
    nparr_b = np.frombuffer(img_b_bytes, np.uint8)
    img_a = cv2.imdecode(nparr_a, cv2.IMREAD_COLOR)
    img_b = cv2.imdecode(nparr_b, cv2.IMREAD_COLOR)

    # 1. 特征检测
    orb = cv2.ORB_create(2000)
    kp1, des1 = orb.detectAndCompute(img_a, None)
    kp2, des2 = orb.detectAndCompute(img_b, None)

    # 2. 特征匹配
    bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=True)
    matches = bf.match(des1, des2)
    matches = sorted(matches, key=lambda x: x.distance)

    # 3. 计算变换矩阵
    if len(matches) > 10:
        src_pts = np.float32([kp1[m.queryIdx].pt for m in matches]).reshape(-1, 1, 2)
        dst_pts = np.float32([kp2[m.trainIdx].pt for m in matches]).reshape(-1, 1, 2)
        M, mask = cv2.findHomography(dst_pts, src_pts, cv2.RANSAC, 5.0)
        h, w, _ = img_a.shape
        img_b_aligned = cv2.warpPerspective(img_b, M, (w, h))
        return img_a, img_b_aligned, True
    else:
        # 特征点不足，退回到手动缩放
        return img_a, img_b, False

def blend_overlay(orig_arr, mask_diff, mask_above, alpha):
    h, w, _ = orig_arr.shape
    overlay = np.full((h, w, 3), 128, dtype=np.uint8)
    overlay[mask_diff & ~mask_above] = [0, 120, 212]
    overlay[mask_above] = [232, 17, 35]
    return (orig_arr.astype(float)*(1-alpha) + overlay.astype(float)*alpha).astype(np.uint8)

# --- UI Application ---

st.set_page_config(page_title="Office-Comparer Pro", layout="wide")
st.markdown(STYLE, unsafe_allow_html=True)
st.markdown(HEADER_HTML, unsafe_allow_html=True)

# Session State
if 'doc_a' not in st.session_state: st.session_state.doc_a = None
if 'doc_b' not in st.session_state: st.session_state.doc_b = None

with st.sidebar:
    st.markdown("## ⚙️ 核心控制器")
    tab_type = st.radio("比对类型", ["Word 文档", "Excel 表格", "图像比对"])
    st.markdown("---")
    threshold = st.slider("差异灵敏度", 0, 255, 30)
    alpha = st.slider("差异覆盖强度", 0, 100, 40) / 100.0
    st.markdown("---")
    if st.button("🗑️ 清除所有缓存"):
        st.cache_data.clear()
        st.session_state.doc_a = st.session_state.doc_b = None
        st.rerun()

# --- Tabs Implementation ---

if tab_type == "Word 文档":
    u1, u2 = st.columns(2)
    fa = u1.file_uploader("上传文档 A", type=["docx"], key="wa")
    fb = u2.file_uploader("上传文档 B", type=["docx"], key="wb")

    if fa and fb:
        lines_a = load_docx_lines(fa.getvalue())
        lines_b = load_docx_lines(fb.getvalue())
        
        if st.session_state.doc_a is None:
            st.session_state.doc_a, st.session_state.doc_b = lines_a, lines_b

        # 批量操作
        c1, c2, c3 = st.columns([2, 2, 6])
        if c1.button("✅ 全部接受 A"): st.session_state.doc_b = list(st.session_state.doc_a); st.rerun()
        if c2.button("✅ 全部接受 B"): st.session_state.doc_a = list(st.session_state.doc_b); st.rerun()
        
        matcher = difflib.SequenceMatcher(None, st.session_state.doc_a, st.session_state.doc_b)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            r = st.columns([0.5, 9, 1, 9, 0.5])
            if tag == 'equal':
                for k in range(i2-i1):
                    r = st.columns([0.5, 9, 1, 9, 0.5])
                    r[0].markdown(f"<div class='line-num'>{i1+k+1}</div>", unsafe_allow_html=True)
                    r[1].markdown(f"<div class='doc-content'>{st.session_state.doc_a[i1+k]}</div>", unsafe_allow_html=True)
                    r[2].markdown("<center style='color:#ccc; line-height:3'>=</center>", unsafe_allow_html=True)
                    r[3].markdown(f"<div class='doc-content'>{st.session_state.doc_b[j1+k]}</div>", unsafe_allow_html=True)
                    r[4].markdown(f"<div class='line-num'>{j1+k+1}</div>", unsafe_allow_html=True)
            else:
                text_a = "\n".join(st.session_state.doc_a[i1:i2])
                text_b = "\n".join(st.session_state.doc_b[j1:j2])
                h_a, h_b = highlight_text_diff(text_a, text_b)
                
                r[0].markdown(f"<div class='line-num'>{i1+1 if i2>i1 else ''}</div>", unsafe_allow_html=True)
                r[1].markdown(f"<div class='doc-content diff-removed'>{h_a if h_a else '(空)'}</div>", unsafe_allow_html=True)
                with r[2]:
                    st.write("") # 间距
                    if st.button("→", key=f"m{i1}{j1}"):
                        st.session_state.doc_b[j1:j2] = st.session_state.doc_a[i1:i2]
                        st.rerun()
                    if st.button("←", key=f"n{i1}{j1}"):
                        st.session_state.doc_a[i1:i2] = st.session_state.doc_b[j1:j2]
                        st.rerun()
                r[3].markdown(f"<div class='doc-content diff-added'>{h_b if h_b else '(空)'}</div>", unsafe_allow_html=True)
                r[4].markdown(f"<div class='line-num'>{j1+1 if j2>j1 else ''}</div>", unsafe_allow_html=True)

        st.divider()
        st.download_button("💾 导出已纠正的 A", save_docx_from_lines(st.session_state.doc_a), "A_fixed.docx")

elif tab_type == "Excel 表格":
    u1, u2 = st.columns(2)
    ea = u1.file_uploader("上传表格 A", type=["xlsx", "xls"], key="ea")
    eb = u2.file_uploader("上传表格 B", type=["xlsx", "xls"], key="eb")

    if ea and eb:
        df_a = load_excel_df(ea.getvalue())
        df_b = load_excel_df(eb.getvalue())
        if df_a is not None and df_b is not None:
            st.info("💡 黄色单元格表示内容不一致")
            styled_a, styled_b = compare_dataframes(df_a, df_b)
            c1, c2 = st.columns(2)
            with c1: st.write("### 文件 A"); st.dataframe(styled_a, use_container_width=True)
            with c2: st.write("### 文件 B"); st.dataframe(styled_b, use_container_width=True)

elif tab_type == "图像比对":
    iu1, iu2 = st.columns(2)
    ia_f = iu1.file_uploader("图片 A", type=["png","jpg","jpeg"], key="ia")
    ib_f = iu2.file_uploader("图片 B", type=["png","jpg","jpeg"], key="ib")

    if ia_f and ib_f:
        with st.spinner("🔍 正在执行 CV2 自动对齐与比对..."):
            a1, a2, aligned = align_images_cv2(ia_f.getvalue(), ib_f.getvalue())
            if aligned: st.success("✅ 已自动识别特征点并完成几何对齐")
            else: st.warning("⚠️ 未能识别足够特征点，仅执行中心对齐")
            
            # 差异计算
            h, w, _ = a1.shape
            md = np.sum(np.abs(a1.astype(np.int16)-a2.astype(np.int16)), 2)>0
            ma = np.mean(np.abs(a1.astype(np.int16)-a2.astype(np.int16)), 2)>threshold
            
            va, vb = blend_overlay(a1, md, ma, alpha), blend_overlay(a2, md, ma, alpha)
            
            MAX_DIM = 1000
            sc = MAX_DIM / max(h, w) if max(h, w) > MAX_DIM else 1
            ns = (int(w*sc), int(h*sc))
            vad = np.array(Image.fromarray(va).resize(ns))
            vbd = np.array(Image.fromarray(vb).resize(ns))
            
            fig = make_subplots(rows=1, cols=2, horizontal_spacing=0.02)
            fig.add_trace(go.Image(z=vad), row=1, col=1)
            fig.add_trace(go.Image(z=vbd), row=1, col=2)
            fig.update_xaxes(matches='x'); fig.update_yaxes(matches='y')
            fig.update_layout(height=700, margin=dict(l=0, r=0, b=0, t=10))
            st.plotly_chart(fig, use_container_width=True)
