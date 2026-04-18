from docx import Document
import pandas as pd
from PIL import Image, ImageDraw
import os

def create_test_data():
    if not os.path.exists("test_data"):
        os.makedirs("test_data")
    
    # 1. Create Word files
    doc1 = Document()
    doc1.add_paragraph("This is the first line.")
    doc1.add_paragraph("This is the second line.")
    doc1.save("test_data/doc_a.docx")
    
    doc2 = Document()
    doc2.add_paragraph("This is the first line.")
    doc2.add_paragraph("This is a changed second line!")
    doc2.save("test_data/doc_b.docx")
    
    # 2. Create Excel files
    df1 = pd.DataFrame({"Name": ["Alice", "Bob"], "Age": [25, 30]})
    df1.to_excel("test_data/excel_a.xlsx", index=False)
    
    df2 = pd.DataFrame({"Name": ["Alice", "Charlie"], "Age": [25, 35]})
    df2.to_excel("test_data/excel_b.xlsx", index=False)
    
    # 3. Create Images
    img1 = Image.new("RGB", (100, 100), color=(255, 255, 255))
    draw1 = ImageDraw.Draw(img1)
    draw1.rectangle([20, 20, 80, 80], fill=(0, 0, 255))
    img1.save("test_data/img_a.png")
    
    img2 = Image.new("RGB", (100, 100), color=(255, 255, 255))
    draw2 = ImageDraw.Draw(img2)
    draw2.rectangle([25, 25, 85, 85], fill=(0, 0, 255)) # Shifted blue box
    img2.save("test_data/img_b.png")

if __name__ == "__main__":
    create_test_data()
    print("Test data created in 'test_data' folder.")
